# player_barograph_streamlit.py
# Run with: streamlit run player_barograph_streamlit.py

import io
import json
from datetime import datetime

import streamlit as st
import plotly.graph_objects as go
import statistics as stats
from docx import Document
from docx.shared import Inches
import re

# ─────────────────────────────────────────────────────────────
# 1. LOAD PLAYER INFO FROM JSON
# ─────────────────────────────────────────────────────────────
with open("Player_Info.json", "r") as f:
    PLAYER_INFO = json.load(f)

# ─────────────────────────────────────────────────────────────
# 2. HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────
def category_average(sub_scores: dict[str, int]) -> float:
    return stats.mean(sub_scores.values())

def build_barograph(
    scores: dict[str, dict[str, int]],
    show_sub: bool,
    show_avg: bool
) -> go.Figure:
    fig = go.Figure()
    categories = list(scores.keys())
    palette = ["#636EFA", "#EF553B", "#00CC96"]

    if show_sub:
        for i, cat in enumerate(categories):
            for j, sub in enumerate(scores[cat]):
                fig.add_trace(go.Bar(
                    x=[i + (j - 1) * 0.22],
                    y=[scores[cat][sub]],
                    width=0.2,
                    name=sub,
                    marker_color=palette[j % len(palette)],
                    legendgroup=sub,
                    showlegend=(i == 0),
                ))

    if show_avg:
        for i, cat in enumerate(categories):
            avg = category_average(scores[cat])
            fig.add_trace(go.Bar(
                x=[i],
                y=[avg],
                width=0.6,
                marker=dict(color="rgba(128,128,128,0.35)"),
                name=f"{cat} Avg",
                legendgroup="avg",
                showlegend=(i == 0),
            ))

    fig.update_layout(
        barmode="overlay",
        title="Skill Barograph",
        xaxis=dict(
            tickmode="array",
            tickvals=list(range(len(categories))),
            ticktext=categories,
            title="Skill Category"
        ),
        yaxis=dict(title="Score (1–120)"),
        legend_title="Legend",
        template="plotly_white"
    )
    return fig

def parse_docx_report(docx_file):
    """
    Parse existing docx for history:
      - category_history: {category: [float, float, ...]}
      - overall_history: list of (date, score)
    """
    doc = Document(docx_file)
    meta = {"player_type": None, "player_name": None, "category_history": {}, "overall_history": []}
    cat_pattern = re.compile(r"^•\s*([^:]+):\s*(.+)$")
    overall_header_seen = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.endswith("Stats Report") and "–" in text:
            p, t = text.replace("Stats Report", "").split("–")
            meta["player_name"] = p.strip()
            meta["player_type"] = t.strip()
        elif cat_pattern.match(text):
            m = cat_pattern.match(text)
            cat = m.group(1).strip()
            vals = [float(x.strip()) for x in m.group(2).split(";") if x.strip()]
            meta["category_history"][cat] = vals
        elif text == "Overall Ratings by Date:":
            overall_header_seen = True
        elif overall_header_seen and re.match(r"^\d{4}-\d{2}-\d{2}:", text):
            date_part, val_part = text.split(":")
            meta["overall_history"].append((date_part.strip(), float(val_part.strip())))
        elif text.startswith("Overall Rating:"):
            # Not used for input, calculated fresh for output
            pass
        elif text == "":
            overall_header_seen = False
    return meta

def create_docx_report(
    player_name: str,
    player_type: str,
    scores: dict[str, dict[str, int]],
    prev_cat_history: dict = None,
    prev_overall_history: list = None,
) -> bytes:
    # Prepare history data for categories
    cat_history = prev_cat_history.copy() if prev_cat_history else {}
    for cat, subdict in scores.items():
        avg = category_average(subdict)
        if cat not in cat_history:
            cat_history[cat] = []
        cat_history[cat].append(round(avg, 1))

    # Prepare overall history
    overall_score = round(stats.mean(category_average(subdict) for subdict in scores.values()), 1)
    overall_history = prev_overall_history.copy() if prev_overall_history else []
    today = datetime.now().strftime("%Y-%m-%d")
    overall_history.append((today, overall_score))

    # Build docx
    fig = build_barograph(scores, show_sub=True, show_avg=True)
    img = fig.to_image(format="png", width=800, height=400)

    doc = Document()
    doc.add_heading(f"{player_name} – {player_type} Stats Report", level=1)
    doc.add_paragraph(f"Report Generated: {today} {datetime.now().strftime('%H:%M:%S')}")

    doc.add_picture(io.BytesIO(img), width=Inches(6))
    doc.add_paragraph()

    # Write category averages/history as a line per category
    for cat in scores.keys():
        vals = cat_history[cat]
        vals_str = "; ".join([f"{v:.1f}" for v in vals])
        doc.add_paragraph(f"• {cat}: {vals_str}", style="List Bullet")

    # Latest overall rating (at the top)
    doc.add_paragraph()
    doc.add_paragraph(f"Overall Rating: {overall_score:.1f} / 120", style="Intense Quote")
    doc.add_paragraph("-" * 38)
    doc.add_paragraph("Overall Ratings by Date:")

    # List all previous overalls, each on its own line
    for date_str, val in overall_history:
        doc.add_paragraph(f"{date_str}: {val:.1f}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────
# 3. STREAMLIT APP (with session state for flow control)
# ─────────────────────────────────────────────────────────────
def main():
    st.title("🏅 Player Stats: Barograph")

    st.sidebar.header("Options")
    show_slider = st.sidebar.checkbox("Show Sliders", True)
    show_num_input = st.sidebar.checkbox("Show Number Inputs", False)
    show_sub = st.sidebar.checkbox("Show Sub-skill Bars", True)
    show_avg = st.sidebar.checkbox("Show Category Average Bars", True)
    if st.sidebar.button("Reset/Start Over"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]

    # SESSION STATE for workflow
    if "mode" not in st.session_state:
        st.session_state.mode = None
        st.session_state.prefill_meta = {}

    if st.session_state.mode is None:
        uploaded = st.file_uploader("Upload .docx report to continue updating, or click below to start fresh.", type="docx")
        create_new = st.button("Create New Report")
        if uploaded is not None:
            st.session_state.mode = "edit"
            st.session_state.prefill_meta = parse_docx_report(uploaded)
        elif create_new:
            st.session_state.mode = "new"
            st.session_state.prefill_meta = {}

    if st.session_state.mode is None:
        st.info("Upload a DOCX to update, or click 'Create New Report' to begin.")
        st.stop()

    prefill_meta = st.session_state.prefill_meta
    use_existing = st.session_state.mode == "edit"

    if use_existing and prefill_meta.get("player_type") and prefill_meta.get("player_name"):
        player_type = prefill_meta["player_type"]
        player_name = prefill_meta["player_name"]
        st.success(f"Loaded: {player_name} – {player_type}")
    else:
        player_type = st.selectbox("Player Type", list(PLAYER_INFO.keys()))
        player_name = st.selectbox("Player Name", list(PLAYER_INFO[player_type]["Players"].keys()))

    st.header(f"{player_name} – {player_type} Skills")

    scores: dict[str, dict[str, int]] = {}
    skillset = PLAYER_INFO[player_type]["Skillset"]

    # Prepare per-category score histories (for docx creation)
    cat_history = prefill_meta.get("category_history", {}) if use_existing else {}
    overall_history = prefill_meta.get("overall_history", []) if use_existing else []

    for cat, subs in skillset.items():
        with st.expander(cat, expanded=True):
            scores[cat] = {}
            if use_existing and cat in cat_history and cat_history[cat]:
                avg_val = int(round(cat_history[cat][-1]))
            else:
                avg_val = 50
            for sub in subs:
                if show_slider and show_num_input:
                    c1, c2 = st.columns(2)
                    val = c1.slider(
                        label=f"{sub}",
                        min_value=1,
                        max_value=100,
                        value=avg_val,
                        key=f"{cat}_{sub}_slider"
                    )
                    val = c2.number_input(
                        label=f"{sub}",
                        min_value=1,
                        max_value=100,
                        value=val,
                        key=f"{cat}_{sub}_input"
                    )
                elif show_slider:
                    val = st.slider(
                        label=sub,
                        min_value=1,
                        max_value=100,
                        value=avg_val,
                        key=f"{cat}_{sub}_slider"
                    )
                else:
                    val = st.number_input(
                        label=sub,
                        min_value=1,
                        max_value=100,
                        value=avg_val,
                        key=f"{cat}_{sub}_input"
                    )
                scores[cat][sub] = val

    specialty = PLAYER_INFO[player_type]["Players"][player_name].get("Specialty", {})
    for spec_cat, subs in specialty.items():
        with st.expander(f"Specialty: {spec_cat}", expanded=True):
            scores[spec_cat] = {}
            if use_existing and spec_cat in cat_history and cat_history[spec_cat]:
                avg_val = int(round(cat_history[spec_cat][-1]))
            else:
                avg_val = 50
            for sub in subs:
                if show_slider and show_num_input:
                    c1, c2 = st.columns(2)
                    val = c1.slider(
                        label=f"{sub}",
                        min_value=1,
                        max_value=100,
                        value=avg_val,
                        key=f"{spec_cat}_{sub}_slider"
                    )
                    val = c2.number_input(
                        label=f"{sub}",
                        min_value=1,
                        max_value=100,
                        value=val,
                        key=f"{spec_cat}_{sub}_input"
                    )
                elif show_slider:
                    val = st.slider(
                        label=sub,
                        min_value=1,
                        max_value=100,
                        value=avg_val,
                        key=f"{spec_cat}_{sub}_slider"
                    )
                else:
                    val = st.number_input(
                        label=sub,
                        min_value=1,
                        max_value=100,
                        value=avg_val,
                        key=f"{spec_cat}_{sub}_input"
                    )
                scores[spec_cat][sub] = val

    fig = build_barograph(scores, show_sub, show_avg)
    st.plotly_chart(fig, use_container_width=True)

    st.markdown("### Category Averages")
    for cat, subdict in scores.items():
        st.write(f"**{cat}**: {category_average(subdict):.1f}")
    overall = stats.mean(category_average(v) for v in scores.values())
    st.markdown(f"**Overall Rating:** {overall:.1f} / 120")

    with st.sidebar:
        if st.button("Save Report as DOCX"):
            docx_bytes = create_docx_report(
                player_name, player_type, scores,
                prev_cat_history=cat_history, prev_overall_history=overall_history
            )
            date_str = datetime.now().strftime("%Y-%m-%d")
            filename = f"{date_str}_{player_name.replace(' ','_')}_{player_type}_Report.docx"
            st.download_button(
                "Download Report",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
